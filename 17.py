from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import calendar


list1=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",1,2,3],[4,5,6,7,8,9,10],[11,12,13,14,15,16,17],[18,19,20,21,22,23,24],[25,26,27,28,29,30,31]]
list2=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28]]
list3=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28],[29,30,31,".",".",".","."]]
list4=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",1,2,3,4],[5,6,7,8,9,10,11],[12,13,14,15,16,17,18],[19,20,21,22,23,24,25],[26,27,28,29,30,".","."]]
list5=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",".",1,2],[3,4,5,6,7,8,9],[10,11,12,13,14,15,16],[17,18,19,20,21,22,23],[24,25,26,27,28,29,30],[31,".",".",".",".",".","."]]
list6=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",1,2,3,4,5,6],[7,8,9,10,11,12,13],[14,15,16,17,18,19,20],[21,22,23,24,25,26,27],[28,29,30,".",".",".","."]]
list7=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",1,2,3,4],[5,6,7,8,9,10,11],[12,13,14,15,16,17,18],[19,20,21,22,23,24,25],[26,27,28,29,30,31,"."]]
list8=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",".",".",1],[2,3,4,5,6,7,8],[9,10,11,12,13,14,15],[16,17,18,19,20,21,22],[23,24,25,26,27,28,29],[30,31,".",".",".",".",","]]
list9=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",1,2,3,4,5],[6,7,8,9,10,11,12],[13,14,15,16,17,18,19],[20,21,22,23,24,25,26],[27,28,29,30,".",".","."]]
list10=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",".",".",1,2,3],[4,5,6,7,8,9,10],[11,12,13,14,15,16,17],[18,19,20,21,22,23,24],[25,26,27,28,29,30,31]]
list11=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[1,2,3,4,5,6,7],[8,9,10,11,12,13,14],[15,16,17,18,19,20,21],[22,23,24,25,26,27,28],[29,30,".",".",".",".","."]]
list12=[["Sun","Mon","Tue","Wed","Thu","Fri","Sat"],[".",".",1,2,3,4,5],[6,7,8,9,10,11,12],[13,14,15,16,17,18,19],[20,21,22,23,24,25,26],[27,28,29,30,31,".","."]]
listfull=[list1,list2,list3,list4,list5,list6,list7,list8,list9,list10,list11,list12]
year="2026"
months=["January","February","March","April","May","June","July","August","September","October","November","December"]
holiday= [{"date": "15-01-2026","name": "Makara Sankranthi","Day": "Thursday"},
        {"date": "26-01-2026", "name": "Republic Day","Day": "Monday"},
        {"date": "19-03-2026", "name": "Chandramana Ugadi", "Day": "Thursday"},
        {"date": "27-03-2026", "name": "Shrirama Navami", "Day": "Friday"},
        {"date": "03-04-2026", "name": "Good Friday", "Day": "Friday"},
        {"date": "14-04-2026", "name": "Ambedkar Jayanthi", "Day": "Tuesday"},
        {"date": "01-05-2026", "name": "Labour day", "Day": "Friday"},
        {"date": "28-05-2026", "name": "Bakrid", "Day": "Thursday"},
        {"date": "26-06-2026", "name": "Moharam", "Day": "Friday"},
        {"date": "15-08-2026", "name": "Independence day", "Day": "Saturday"},
        {"date": "28-01-2026", "name": "Ed-Milad", "Day": "wednesday"},
        {"date": "05-09-2026", "name": "Vitla pindi", "Day": "Saturday"},
        {"date": "14-09-2026", "name": "Ganesha chathurthi", "Day": "Monday"},
        {"date": "02-10-2026", "name": "Gandi Jayanthi", "Day": "Friday"},
        {"date": "20-10-2026", "name": "Maha Navami / Ayudha pooja", "Day": "Tuesday"},
        {"date": "21-10-2026", "name": "Dassehra / vijayadashami", "Day": "wednesday"},
        {"date": "10-11-2026", "name": "Deepavali", "Day": "Tuesday"},
        {"date": "25-12-2026", "name": "Christmas", "Day": "Friday"}]
name1="Dhyan Kulal"
doc=Document()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Calendar 2026")
run.font.color.rgb = RGBColor(255, 0, 0)
run.font.size = Pt(20)

year = 2026
cal = calendar.TextCalendar(calendar.SUNDAY)
table4 = doc.add_table(rows=4, cols=3)
month_no = 1
for row in range(4):
    for col in range(3):
        cell = table4.cell(row, col)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(cal.formatmonth(year, month_no))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        month_no += 1
run = doc.add_paragraph().add_run(name1)
run.font.color.rgb = RGBColor(200, 200, 200)
doc.add_page_break()
doc.add_paragraph()
count = 0
for i in range(len(listfull)):
    doc.add_picture("tiger.jpg")
    p=doc.add_paragraph().add_run("2026")
    p.bold = True
    p = doc.add_paragraph("")
    p.add_run(" "*100)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(months[i]))
    run.bold = True
    run.font.size = Pt(14)

    rows = len(listfull[i])
    table = doc.add_table(rows=rows, cols=7)
    table.style = "Table Grid"   
    
    for j in range(rows):
        for k in range(len(listfull[i][j])):
            table.rows[j].height =Inches(0.9)
            table.columns[j].width =Inches(0.8)
            #table.cell(j, k).paragraphs(str(listfull[i][j][k]))
            cell = table.cell(j, k)
    
            cell.text = str(listfull[i][j][k])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if k == 0:
                run.font.color.rgb = RGBColor(255, 55, 0)
    run = doc.add_paragraph().add_run(name1)
    run.font.color.rgb = RGBColor(200, 200, 200)
    doc.add_page_break()

doc.add_picture("tiger.jpg")
doc.add_heading("HOLIDAYS")
table2=doc.add_table(rows=18, cols=3)
for i in range(len(holiday)):
    h = holiday[i]
    run=table2.cell(i,0)
    run.text=h["date"]
for i in range(len(holiday)):
    h=holiday[i]
    run=table2.cell(i,1)
    run.text=h["Day"]
for i in range(len(holiday)):
    h=holiday[i]
    run=table2.cell(i,2)
    run.text=h["name"]
run = doc.add_paragraph().add_run(name1)
run.font.color.rgb = RGBColor(200, 200, 200)
doc.save("Calendar2026.docx")